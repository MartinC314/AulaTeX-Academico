from __future__ import annotations

from pathlib import Path
import json
from datetime import datetime
import pandas as pd
from docx import Document
from docx.shared import Pt

from .fichas import ConceptFicha
from .planeacion_parser import PlaneacionAnalizada, planeacion_to_dict


def fichas_to_rows(fichas: list[ConceptFicha]) -> list[dict]:
    rows: list[dict] = []
    for i, ficha in enumerate(fichas, start=1):
        if not ficha.hits:
            rows.append({
                "ficha": i,
                "concepto": ficha.concept,
                "calidad_ficha": ficha.quality_label,
                "mejor_similitud": round(ficha.best_score, 4),
                "similitud_promedio": round(ficha.average_score, 4),
                "fuente": "",
                "tipo_fuente": "",
                "ubicacion": "",
                "pagina_o_bloque": "",
                "similitud": "",
                "detalle_similitud": "",
                "cita_textual": "",
                "fragmento_id": "",
                "ruta_fuente": "",
                "ubicaciones_agrupadas": ficha.locations_text,
                "observacion": ficha.observation,
            })
            continue
        for hit in ficha.hits:
            rows.append({
                "ficha": i,
                "concepto": ficha.concept,
                "calidad_ficha": ficha.quality_label,
                "mejor_similitud": round(ficha.best_score, 4),
                "similitud_promedio": round(ficha.average_score, 4),
                "fuente": hit.source_name,
                "tipo_fuente": hit.source_type,
                "ubicacion": hit.location,
                "pagina_o_bloque": hit.page,
                "similitud": round(hit.score, 4),
                "detalle_similitud": hit.score_details,
                "cita_textual": hit.quote,
                "fragmento_id": hit.fragment_id,
                "ruta_fuente": hit.source_path,
                "ubicaciones_agrupadas": ficha.locations_text,
                "observacion": ficha.observation,
            })
    return rows


def export_markdown(fichas: list[ConceptFicha], output_path: Path) -> None:
    lines: list[str] = ["# Fichas de conceptos", ""]
    lines.append("Las citas textuales se extrajeron de las fuentes de entrada. La búsqueda puede agrupar referencias del mismo concepto aunque estén en archivos distintos.")
    lines.append("")
    for i, ficha in enumerate(fichas, start=1):
        lines.append(f"## Ficha {i:02d}. {ficha.concept}")
        lines.append("")
        lines.append(f"**Calidad estimada:** {ficha.quality_label}")
        lines.append(f"**Mejor similitud:** {ficha.best_score:.4f}")
        lines.append(f"**Similitud promedio:** {ficha.average_score:.4f}")
        lines.append(f"**Fuentes:** {ficha.sources_text}")
        lines.append(f"**Ubicaciones:** {ficha.locations_text}")
        lines.append("")
        if ficha.hits:
            for hit in ficha.hits:
                lines.append(f"### {hit.source_name} — {hit.location} — similitud {hit.score:.4f}")
                lines.append("")
                if hit.score_details:
                    lines.append(f"**Detalle de puntaje:** {hit.score_details}")
                    lines.append("")
                lines.append(f"> {hit.quote}")
                lines.append("")
        else:
            lines.append("No se localizaron citas textuales por encima del umbral configurado.")
            lines.append("")
        lines.append(f"**Observación automática:** {ficha.observation}")
        lines.append("")
    output_path.write_text("\n".join(lines), encoding="utf-8")


def export_excel(rows: list[dict], output_path: Path) -> None:
    df = pd.DataFrame(rows)
    with pd.ExcelWriter(output_path, engine="openpyxl") as writer:
        df.to_excel(writer, index=False, sheet_name="fichas")
        ws = writer.book["fichas"]
        ws.freeze_panes = "A2"
        widths = {"A": 10, "B": 34, "C": 34, "D": 14, "E": 16, "F": 16, "G": 12, "H": 90, "I": 34, "J": 55, "K": 70, "L": 70}
        for col, width in widths.items():
            ws.column_dimensions[col].width = width
        for row in ws.iter_rows(min_row=2):
            for cell in row:
                cell.alignment = cell.alignment.copy(wrap_text=True, vertical="top")


def export_csv(rows: list[dict], output_path: Path) -> None:
    pd.DataFrame(rows).to_csv(output_path, index=False, encoding="utf-8-sig")


def export_json(rows: list[dict], output_path: Path) -> None:
    output_path.write_text(json.dumps(rows, ensure_ascii=False, indent=2), encoding="utf-8")


def export_docx(fichas: list[ConceptFicha], output_path: Path) -> Path:
    doc = Document()
    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10)
    doc.add_heading("Fichas de conceptos", level=1)
    doc.add_paragraph("Archivo generado automáticamente. Las citas textuales se extrajeron de las fuentes de entrada.")

    for i, ficha in enumerate(fichas, start=1):
        doc.add_heading(f"Ficha {i:02d}. {ficha.concept}", level=2)
        p = doc.add_paragraph()
        p.add_run("Fuentes: ").bold = True
        p.add_run(ficha.sources_text)
        p = doc.add_paragraph()
        p.add_run("Ubicaciones: ").bold = True
        p.add_run(ficha.locations_text)

        if ficha.hits:
            for hit in ficha.hits:
                doc.add_heading(f"{hit.source_name} — {hit.location} — similitud {hit.score:.4f}", level=3)
                q = doc.add_paragraph()
                q.paragraph_format.left_indent = Pt(18)
                q.add_run(f"“{hit.quote}”")
        else:
            doc.add_paragraph("No se localizaron citas textuales por encima del umbral configurado.")

        p = doc.add_paragraph()
        p.add_run("Observación automática: ").bold = True
        p.add_run(ficha.observation)

    try:
        doc.save(output_path)
        return output_path
    except PermissionError:
        stamp = datetime.now().strftime("%Y%m%d-%H%M%S")
        fallback = output_path.with_name(f"{output_path.stem}-{stamp}{output_path.suffix}")
        doc.save(fallback)
        return fallback


def _build_idea_rows(fichas: list[ConceptFicha]) -> list[dict]:
    ideas: list[dict] = []
    for ficha in fichas:
        if not ficha.hits:
            continue
        hit = ficha.hits[0]
        ideas.append({
            "concepto": ficha.concept,
            "tipo": "idea",
            "idea_base": hit.quote,
            "fuente": hit.source_name,
            "ubicacion": hit.location,
            "funcion_sugerida": "desarrollo",
            "observacion": ficha.observation,
        })
    return ideas


def _build_traceability(rows: list[dict]) -> list[dict]:
    trace: list[dict] = []
    for row in rows:
        trace.append({
            "concepto": row.get("concepto", ""),
            "fuente": row.get("fuente", ""),
            "ruta_fuente": row.get("ruta_fuente", ""),
            "ubicacion": row.get("ubicacion", ""),
            "fragmento_id": row.get("fragmento_id", ""),
            "cita_textual": row.get("cita_textual", ""),
        })
    return trace


def export_all(
    fichas: list[ConceptFicha],
    output_dir: str | Path,
    *,
    planeacion_analysis: PlaneacionAnalizada | None = None,
    planeacion_remote: dict | None = None,
    planeacion_final: dict | None = None,
    subject_slug: str | None = None,
    concept_curation: dict | None = None,
    refinement_diagnostics: dict | None = None,
    conceptos: list[str] | None = None,
) -> dict[str, Path]:
    out = Path(output_dir)
    out.mkdir(parents=True, exist_ok=True)
    rows = fichas_to_rows(fichas)
    paths = {
        "markdown": out / "fichas_conceptos.md",
        "excel": out / "fichas_conceptos.xlsx",
        "csv": out / "fichas_conceptos.csv",
        "json": out / "fichas_conceptos.json",
        "word": out / "fichas_conceptos.docx",
        "planeacion": out / "resumen_planeacion.json",
        "planeacion_local": out / "resumen_planeacion_local.json",
        "planeacion_anthropic": out / "resumen_planeacion_anthropic.json",
        "depuracion": out / "depuracion_conceptos.json",
        "refinamiento": out / "refinamiento_fichas.json",
        "conceptos": out / "conceptos_detectados.json",
        "ideas": out / "ideas_detectadas.json",
        "trazabilidad": out / "trazabilidad_fuentes.json",
    }
    export_markdown(fichas, paths["markdown"])
    export_excel(rows, paths["excel"])
    export_csv(rows, paths["csv"])
    export_json(rows, paths["json"])
    paths["word"] = export_docx(fichas, paths["word"])
    paths["conceptos"].write_text(json.dumps(conceptos or [f.concept for f in fichas], ensure_ascii=False, indent=2), encoding="utf-8")
    paths["ideas"].write_text(json.dumps(_build_idea_rows(fichas), ensure_ascii=False, indent=2), encoding="utf-8")
    paths["trazabilidad"].write_text(json.dumps(_build_traceability(rows), ensure_ascii=False, indent=2), encoding="utf-8")
    if concept_curation is not None:
        payload = {"subject_slug": subject_slug, **concept_curation}
        paths["depuracion"].write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
    if refinement_diagnostics is not None:
        paths["refinamiento"].write_text(json.dumps(refinement_diagnostics, ensure_ascii=False, indent=2), encoding="utf-8")
    if planeacion_analysis is not None:
        local_dict = planeacion_to_dict(planeacion_analysis)
        paths["planeacion_local"].write_text(json.dumps(local_dict, ensure_ascii=False, indent=2), encoding="utf-8")
        final_dict = planeacion_final or local_dict
        paths["planeacion"].write_text(json.dumps(final_dict, ensure_ascii=False, indent=2), encoding="utf-8")
    if planeacion_remote is not None:
        paths["planeacion_anthropic"].write_text(json.dumps(planeacion_remote, ensure_ascii=False, indent=2), encoding="utf-8")
    return paths
