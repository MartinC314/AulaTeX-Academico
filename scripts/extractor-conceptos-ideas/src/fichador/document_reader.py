from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
import re

from docx import Document

from .pdf_reader import PageText, extract_pages, make_source_id
from .pandoc_utils import normalize_text_with_pandoc, pandoc_available, pdf_pandoc_normalization_enabled
from .preprocessing import normalize_spaces


SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".md", ".markdown"}


@dataclass(frozen=True)
class SourceLoadReport:
    loaded_files: list[Path]
    skipped_files: list[Path]
    pages_or_blocks: int


def parse_extensions(value: str | None) -> set[str]:
    if not value:
        return set(SUPPORTED_EXTENSIONS)
    result: set[str] = set()
    for part in value.split(","):
        part = part.strip().lower()
        if not part:
            continue
        if not part.startswith("."):
            part = f".{part}"
        result.add(part)
    return result or set(SUPPORTED_EXTENSIONS)


def discover_source_files(path: str | Path, *, recursive: bool = False, extensions: set[str] | None = None) -> tuple[list[Path], list[Path]]:
    root = Path(path)
    if not root.exists():
        raise FileNotFoundError(f"No existe la ruta de fuentes: {root}")

    allowed = extensions or set(SUPPORTED_EXTENSIONS)
    if root.is_file():
        if root.suffix.lower() in allowed:
            return [root], []
        return [], [root]

    pattern = "**/*" if recursive else "*"
    files: list[Path] = []
    skipped: list[Path] = []
    for item in sorted(root.glob(pattern)):
        if not item.is_file():
            continue
        if item.name.startswith("~$"):
            continue
        if item.suffix.lower() in allowed:
            files.append(item)
        else:
            skipped.append(item)
    return files, skipped


def _read_text_file(path: Path) -> str:
    encodings = ["utf-8-sig", "utf-8", "cp1252", "latin-1"]
    last_error: Exception | None = None
    for enc in encodings:
        try:
            return path.read_text(encoding=enc)
        except UnicodeDecodeError as exc:
            last_error = exc
    if last_error:
        raise last_error
    return path.read_text()


def _split_long_text(text: str, *, max_chars: int = 6000) -> list[str]:
    text = normalize_spaces(text)
    if not text:
        return []
    paragraphs = [normalize_spaces(p) for p in re.split(r"\n\s*\n+", text) if normalize_spaces(p)]
    if not paragraphs:
        paragraphs = [text]
    blocks: list[str] = []
    buffer = ""
    for paragraph in paragraphs:
        if len(paragraph) > max_chars:
            if buffer:
                blocks.append(buffer.strip())
                buffer = ""
            for start in range(0, len(paragraph), max_chars):
                chunk = paragraph[start:start + max_chars].strip()
                if chunk:
                    blocks.append(chunk)
            continue
        candidate = f"{buffer}\n\n{paragraph}".strip() if buffer else paragraph
        if len(candidate) <= max_chars:
            buffer = candidate
        else:
            if buffer:
                blocks.append(buffer.strip())
            buffer = paragraph
    if buffer:
        blocks.append(buffer.strip())
    return blocks


def extract_docx_blocks(path: str | Path) -> list[PageText]:
    p = Path(path)
    source_id = make_source_id(p)
    doc = Document(p)
    pieces: list[str] = []

    for paragraph in doc.paragraphs:
        text = normalize_spaces(paragraph.text)
        if text:
            pieces.append(text)

    for table in doc.tables:
        for row in table.rows:
            cells = [normalize_spaces(cell.text) for cell in row.cells if normalize_spaces(cell.text)]
            if cells:
                pieces.append(" | ".join(cells))

    blocks = _split_long_text("\n\n".join(pieces))
    return [
        PageText(
            page=i,
            text=block,
            source_id=source_id,
            source_name=p.name,
            source_path=str(p),
            source_type="docx",
            location_label=f"bloque {i}",
        )
        for i, block in enumerate(blocks, start=1)
    ]


def extract_plain_text_blocks(path: str | Path) -> list[PageText]:
    p = Path(path)
    source_id = make_source_id(p)
    text = _read_text_file(p)
    blocks = _split_long_text(text)
    source_type = p.suffix.lower().lstrip(".") or "txt"
    return [
        PageText(
            page=i,
            text=block,
            source_id=source_id,
            source_name=p.name,
            source_path=str(p),
            source_type=source_type,
            location_label=f"bloque {i}",
        )
        for i, block in enumerate(blocks, start=1)
    ]


def extract_pages_from_file(path: str | Path) -> list[PageText]:
    p = Path(path)
    suffix = p.suffix.lower()
    if suffix == ".pdf":
        return extract_pages(p)
    if suffix == ".docx":
        return extract_docx_blocks(p)
    if suffix in {".txt", ".md", ".markdown"}:
        return extract_plain_text_blocks(p)
    raise ValueError(f"Tipo de fuente no soportado: {p}")


def extract_pages_from_sources(source_path: str | Path, *, recursive: bool = False, extensions: set[str] | None = None) -> tuple[list[PageText], SourceLoadReport]:
    files, skipped = discover_source_files(source_path, recursive=recursive, extensions=extensions)
    pages: list[PageText] = []
    loaded: list[Path] = []
    failed: list[Path] = []

    for file in files:
        try:
            file_pages = extract_pages_from_file(file)
        except Exception:
            failed.append(file)
            continue
        if any(p.text.strip() for p in file_pages):
            pages.extend(file_pages)
            loaded.append(file)
        else:
            failed.append(file)

    return pages, SourceLoadReport(loaded_files=loaded, skipped_files=skipped + failed, pages_or_blocks=len(pages))


def read_any_text_file(path: str | Path) -> str:
    """Lee una planeación o archivo de conceptos como texto, sin convertirlo en corpus."""
    p = Path(path)
    if not p.exists() or not p.is_file():
        return ""
    suffix = p.suffix.lower()
    try:
        if suffix == ".pdf":
            text = "\n\n".join(page.text for page in extract_pages(p))
            if text.strip() and pandoc_available() and pdf_pandoc_normalization_enabled():
                text = normalize_text_with_pandoc(text, from_format="markdown", to_format="plain")
            return text
        if suffix == ".docx":
            return "\n\n".join(block.text for block in extract_docx_blocks(p))
        if suffix in {".txt", ".md", ".markdown"}:
            return _read_text_file(p)
    except Exception:
        return ""
    return ""
