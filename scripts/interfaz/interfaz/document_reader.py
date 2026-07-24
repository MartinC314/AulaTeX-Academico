from __future__ import annotations

from pathlib import Path


SUPPORTED_DOCUMENT_SUFFIXES = {".txt", ".md", ".pdf", ".docx"}


def _ensure_non_empty_text(text: str, empty_message: str = "El documento está vacío.") -> str:
    cleaned = text.strip()
    if not cleaned:
        raise RuntimeError(empty_message)
    return cleaned


def _read_plain_text(path: Path) -> str:
    return _ensure_non_empty_text(path.read_text(encoding="utf-8", errors="ignore"))


def _read_pdf_text(path: Path) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as exc:  # pragma: no cover - depende del entorno
        raise RuntimeError("Falta instalar pypdf para leer archivos PDF.") from exc

    reader = PdfReader(str(path))
    pages = [page.extract_text() or "" for page in reader.pages]
    return _ensure_non_empty_text(
        "\n\n".join(pages),
        "No pude extraer texto del PDF. Si es un PDF escaneado, conviértelo con OCR antes de subirlo.",
    )


def _read_docx_text(path: Path) -> str:
    try:
        from docx import Document
    except ImportError as exc:  # pragma: no cover - depende del entorno
        raise RuntimeError("Falta instalar python-docx para leer archivos DOCX.") from exc

    document = Document(str(path))
    paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
    table_cells = []
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text:
                    table_cells.append(text)

    parts = paragraphs + table_cells
    return _ensure_non_empty_text("\n\n".join(parts))


def read_document_text(file_path: str | Path) -> str:
    path = Path(file_path)
    if not path.exists() or not path.is_file():
        raise RuntimeError(f"No existe el archivo: {path}")

    suffix = path.suffix.lower()
    if suffix in {".txt", ".md"}:
        return _read_plain_text(path)
    if suffix == ".pdf":
        return _read_pdf_text(path)
    if suffix == ".docx":
        return _read_docx_text(path)

    raise RuntimeError(
        f"Formato no soportado: {suffix}. Usa PDF, TXT, MD o DOCX."
    )
